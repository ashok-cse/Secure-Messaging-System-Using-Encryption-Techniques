"""
docx_to_pdf.py
--------------
Lightweight DOCX -> PDF converter built specifically for the project report.
It reads the document in body order (paragraphs and tables) with python-docx
and renders a PDF with reportlab, preserving headings, bullet/number lists,
tables, inline images and the embedded screenshots.

Usage: python docx_to_pdf.py <input.docx> <output.pdf>
"""

import io
import sys
from xml.sax.saxutils import escape

import docx
from docx.oxml.ns import qn
from PIL import Image as PILImage
from reportlab.lib.enums import TA_JUSTIFY, TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.lib import colors
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Image, Table, TableStyle, PageBreak,
)

PAGE_W, PAGE_H = A4
MARGIN = 2 * cm
CONTENT_W = PAGE_W - 2 * MARGIN

styles = getSampleStyleSheet()
S = {
    "Title": ParagraphStyle("rTitle", parent=styles["Normal"], fontName="Helvetica-Bold",
                            fontSize=20, leading=24, alignment=TA_CENTER, spaceAfter=6),
    "Cover": ParagraphStyle("rCover", parent=styles["Normal"], fontSize=12,
                            leading=16, alignment=TA_CENTER, spaceAfter=4),
    "CoverBold": ParagraphStyle("rCoverB", parent=styles["Normal"], fontName="Helvetica-Bold",
                                fontSize=14, leading=18, alignment=TA_CENTER, spaceAfter=4),
    "H1": ParagraphStyle("rH1", parent=styles["Normal"], fontName="Helvetica-Bold",
                        fontSize=15, leading=19, spaceBefore=14, spaceAfter=6,
                        textColor=colors.HexColor("#1a3c6e")),
    "H2": ParagraphStyle("rH2", parent=styles["Normal"], fontName="Helvetica-Bold",
                        fontSize=12.5, leading=16, spaceBefore=10, spaceAfter=4,
                        textColor=colors.HexColor("#2e5a8a")),
    "Body": ParagraphStyle("rBody", parent=styles["Normal"], fontSize=10.5, leading=15,
                          alignment=TA_JUSTIFY, spaceAfter=6),
    "Bullet": ParagraphStyle("rBullet", parent=styles["Normal"], fontSize=10.5, leading=15,
                            leftIndent=18, bulletIndent=6, spaceAfter=2),
    "Number": ParagraphStyle("rNumber", parent=styles["Normal"], fontSize=10.5, leading=15,
                            leftIndent=18, bulletIndent=6, spaceAfter=2),
    "Code": ParagraphStyle("rCode", parent=styles["Normal"], fontName="Courier",
                          fontSize=8.5, leading=11, spaceAfter=6, textColor=colors.HexColor("#222222")),
    "CellHdr": ParagraphStyle("rCellH", parent=styles["Normal"], fontName="Helvetica-Bold",
                            fontSize=9.5, leading=12, textColor=colors.white),
    "Cell": ParagraphStyle("rCell", parent=styles["Normal"], fontSize=9.5, leading=12),
}


def rt(text):
    """Escape text for reportlab paragraph markup and convert breaks."""
    return escape(text).replace("\n", "<br/>")


def images_in(element, part):
    """Return list of image blobs (bytes) referenced inside an XML element."""
    blobs = []
    for blip in element.iter(qn("a:blip")):
        rid = blip.get(qn("r:embed"))
        if rid and rid in part.related_parts:
            blobs.append(part.related_parts[rid].blob)
    return blobs


def scaled_image(blob, max_w):
    """Return a reportlab Image scaled to fit max_w, preserving aspect ratio."""
    bio = io.BytesIO(blob)
    px_w, px_h = PILImage.open(bio).size
    bio.seek(0)
    w_pt = px_w * 72.0 / 96.0
    h_pt = px_h * 72.0 / 96.0
    if w_pt > max_w:
        scale = max_w / w_pt
        w_pt, h_pt = w_pt * scale, h_pt * scale
    img = Image(bio, width=w_pt, height=h_pt)
    img.hAlign = "CENTER"
    return img


def is_code_para(p):
    if "\n" in p.text:
        return True
    for r in p.runs:
        if (r.font.name or "").lower() in ("consolas", "courier new", "courier"):
            return True
    return False


def convert(in_path, out_path):
    d = docx.Document(in_path)
    part = d.part
    para_map = {p._element: p for p in d.paragraphs}
    tbl_map = {t._element: t for t in d.tables}

    story = []
    in_cover = True
    number_counter = 0

    for child in d.element.body.iterchildren():
        if child in para_map:
            p = para_map[child]
            style_name = p.style.name if p.style else "Normal"
            text = p.text.strip()

            # Inline images in a paragraph
            blobs = images_in(p._element, part)
            if blobs:
                for b in blobs:
                    story.append(scaled_image(b, CONTENT_W))
                    story.append(Spacer(1, 6))
                continue

            if style_name == "Heading 1":
                in_cover = False
                number_counter = 0
                story.append(Paragraph(rt(text), S["H1"]))
            elif style_name == "Heading 2":
                number_counter = 0
                story.append(Paragraph(rt(text), S["H2"]))
            elif style_name == "List Bullet":
                story.append(Paragraph(rt(text), S["Bullet"], bulletText="•"))
            elif style_name == "List Number":
                number_counter += 1
                story.append(Paragraph(rt(text), S["Number"], bulletText=f"{number_counter}."))
            else:
                if not text:
                    story.append(Spacer(1, 4))
                    continue
                number_counter = 0
                if in_cover:
                    # First line big, the two title lines bold, rest centered
                    if text == "SOFTWARE SECURITY":
                        story.append(Spacer(1, 30))
                        story.append(Paragraph(rt(text), S["Title"]))
                    elif "Secure Messaging System" in text or "Encryption Techniques" in text:
                        story.append(Paragraph(rt(text), S["CoverBold"]))
                    else:
                        story.append(Paragraph(rt(text), S["Cover"]))
                elif is_code_para(p):
                    story.append(Paragraph(rt(p.text), S["Code"]))
                else:
                    story.append(Paragraph(rt(text), S["Body"]))

        elif child in tbl_map:
            t = tbl_map[child]
            data = []
            for ri, row in enumerate(t.rows):
                cells = []
                for cell in row.cells:
                    cblobs = images_in(cell._element, part)
                    if cblobs:
                        # full-width single-cell screenshot tables
                        cells.append(scaled_image(cblobs[0], CONTENT_W - 12))
                    else:
                        cstyle = S["CellHdr"] if ri == 0 else S["Cell"]
                        cells.append(Paragraph(rt(cell.text.strip()), cstyle))
                data.append(cells)

            ncols = len(t.columns)
            col_w = CONTENT_W / ncols
            tbl = Table(data, colWidths=[col_w] * ncols)
            tstyle = [
                ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#999999")),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("LEFTPADDING", (0, 0), (-1, -1), 5),
                ("RIGHTPADDING", (0, 0), (-1, -1), 5),
                ("TOPPADDING", (0, 0), (-1, -1), 4),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
            ]
            # header shading only for multi-column data tables
            if ncols > 1:
                tstyle.append(("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#2e5a8a")))
            tbl.setStyle(TableStyle(tstyle))
            story.append(Spacer(1, 4))
            story.append(tbl)
            story.append(Spacer(1, 8))

        # add a page break right after the cover block
        if in_cover is False and len(story) and isinstance(story[-1], Paragraph) \
                and story[-1].style.name == "rH1" and "ABSTRACT" in story[-1].text:
            # ensure abstract starts on a new page
            idx = len(story) - 1
            story.insert(idx, PageBreak())
            in_cover = None  # only once

    doc = SimpleDocTemplate(out_path, pagesize=A4,
                            leftMargin=MARGIN, rightMargin=MARGIN,
                            topMargin=MARGIN, bottomMargin=MARGIN,
                            title="Secure Messaging System Using Encryption Techniques")
    doc.build(story)
    print("PDF written:", out_path)


if __name__ == "__main__":
    convert(sys.argv[1], sys.argv[2])
